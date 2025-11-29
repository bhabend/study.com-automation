from docx import Document
import difflib
import os
from urllib.parse import urlparse
from app.utils import log


def slugify_url(url: str) -> str:
    """
    Converts: https://study.com/abc/page.html
    Into: study.com-abc-page
    """
    parsed = urlparse(url)
    path = parsed.path.replace("/", "-").strip("-")
    host = parsed.netloc.replace("www.", "")
    return f"{host}{('-' + path if path else '')}"


def ensure_output_folder():
    os.makedirs("outputs", exist_ok=True)


def save_docx(path, title, body_text):
    doc = Document()
    doc.add_heading(title, level=1)

    for p in body_text.split("\n"):
        if p.strip():
            doc.add_paragraph(p)
        else:
            doc.add_paragraph("")  # blank line

    doc.save(path)
    log(f"Saved DOCX → {path}")


def generate_diff(original_text, optimized_text):
    """
    Returns a unified diff string.
    """
    orig_lines = original_text.splitlines()
    opt_lines = optimized_text.splitlines()

    diff = difflib.unified_diff(
        orig_lines,
        opt_lines,
        fromfile="original",
        tofile="optimized",
        lineterm=""
    )
    return "\n".join(diff)


# ---------------------------------------------------------
# MAIN EXPORT FUNCTIONS CALLED BY STREAMLIT UI
# ---------------------------------------------------------

def export_original(url: str, original_text: str):
    ensure_output_folder()
    slug = slugify_url(url)
    path = f"outputs/{slug}_original.docx"
    save_docx(path, "Original Content", original_text)
    return path


def export_rewritten(url: str, rewritten_text: str):
    ensure_output_folder()
    slug = slugify_url(url)
    path = f"outputs/{slug}_rewritten.docx"
    save_docx(path, "Rewritten Content", rewritten_text)
    return path


def export_diff(url: str, original_text: str, rewritten_text: str):
    ensure_output_folder()
    slug = slugify_url(url)

    diff_text = generate_diff(original_text, rewritten_text)

    # Save diff as a Word file for consistent export
    path = f"outputs/{slug}_diff.docx"
    save_docx(path, "Before/After Diff", diff_text)
    return path
